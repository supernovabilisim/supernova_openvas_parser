from module.translator import Translator

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document

import re


class FileOps:

    def __init__(self, report_location: str, document_location: str,
                 nvt_options: list):

        # Setting regex for ip addresses
        self.ip_re = r'^((25[0-5]|(2[0-4]|1[0-9]|[1-9]|)[0-9])(\.(?!$)|$)){4}$'

        # Setting regex for CVE's
        self.cve_re = r'^CVE-\d{4}-\d{4,7}$'

        # Setting regex for urls
        self.url_re = r'(https?:\/\/(?:www\.|(?!www))[a-zA-Z0-9][a-zA-Z0-9-]+\
        [a-zA-Z0-9]\.[^\s]{2,}|www\.[a-zA-Z0-9][a-zA-Z0-9-]+[a-zA-Z0-9]\.[^\s]\
        {2,}|https?:\/\/(?:www\.|(?!www))[a-zA-Z0-9]+\.[^\s]{2,}|www\.\
        [a-zA-Z0-9]+\.[^\s]{2,})'

        # Set nvt_options
        self.nvt_options = nvt_options

        # Translator stuff
        self.translate_list = ["Impact", "Solution", "Threat"]
        self.translator = Translator()

        # Create dicts for checking
        self.ntp_dict = {"NVT": 0, "Threat": 0, "Port": 0}
        self.is_dict = {"Impact": 0, "Solution": 0}

        # Create and set document stuff
        self.data = []
        self.doc = Document()
        self.document_location = document_location
        self.doc.add_heading('Report', 0)

        # Report file location
        self.report_location = report_location

    def read_file(self) -> list:

        file = open(self.report_location, 'r')
        contents = file.readlines()

        # Appending new lines
        for _ in range(100):
            contents.append("\n")

        return contents

    def replace_and_match(self, line: str,
                          string_to_replace: str) -> re.Match[str]:

        line = line.replace('\n', '')
        line = line.replace(string_to_replace, '')
        match = re.match(self.ip_re, line)

        return match

    def add_to_document(self, first_column: str, second_column: str) -> None:

        # Deleting \n
        first_column = first_column.replace(':\n', '')

        if first_column != "References":
            second_column = second_column.replace('!\n', '').replace('\n ', '')

        if first_column in self.translate_list:
            first_column = self.translator.translate(first_column)
            second_column = self.translator.translate(second_column)

        # Delete whitespace
        first_column = first_column.lstrip(' ')
        second_column = second_column.lstrip(' ').rstrip()

        # Appending to the document list
        self.data.append(
            [first_column, second_column]
        )

    def write_multiple_lines(self, from_line: str,
                             index: str, file: list) -> None:

        # Create a string for a new column
        new_column = ""
        new_index = index+30

        # Get the explanation of the dicts items
        for num in range(index+1, new_index):

            if file[num] == "\n":
                self.add_to_document(
                    first_column=from_line,
                    second_column=new_column
                )
                break

            new_column += file[num]

    def find_references(self, index: str, file: list):

        i = 0
        new_column = ""

        for line in range(index, index+70):

            if "NVT" in file[line]:
                break

            if i < 3:
                length = len(new_column)
                new = file[line].lstrip()
                new = new.replace('url: ', '')
                new = new.replace('cve: ', '')
                url_re = re.match(self.url_re, new)
                cve_re = re.match(self.cve_re, new)

                if url_re is not None:
                    new_column += url_re.group() + "\n"

                if cve_re is not None:
                    new_column += cve_re.group() + "\n"

                if len(new_column) != length:
                    i += 1

            else:
                break

        if new_column != "":
            self.add_to_document(
                    first_column='References',
                    second_column=new_column
                )

    def reset_dicts(self) -> None:

        for i in self.ntp_dict:
            self.ntp_dict[i] = 0

        for i in self.is_dict:
            self.is_dict[i] = 0

    def save_file(self) -> str:

        new_data = []
        [print(i) for i in self.data]

        if len(self.nvt_options) != 3:
            for index, row in enumerate(self.data):

                if "Host" in row[0]:
                    new_data.append(row)

                if "NVT" in row[0]:
                    
                    data = self.data[index+2][1].split(' ')[0]

                    if self.data[index+1][1].split(' ')[0] in self.nvt_options:
                        new_data.append(row)
 
                    elif data in self.nvt_options:
                        new_data.append(row)

                if "Tehdit" in row[0]:
                    if row[1].split(' ')[0] in self.nvt_options:
                        for i in range(index, index+5):

                            s = self.data[i][0]

                            if s == "Host" or s == "NVT":
                                break

                            else:
                                new_data.append(self.data[i])

        if new_data != []:
            self.data = new_data.copy()

        for col1, col2 in self.data:

            if col1 == "Host":
                self.doc.add_heading(f"\n{col1}: {col2}", level=1)
                continue

            if "NVT" in col1:
                self.doc.add_heading(f"{col1}: {col2}", level=2)
                table = self.doc.add_table(rows=0, cols=2)
                continue

            row = table.add_row().cells

            row[0].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row[0].text = col1

            if col2 == "Çözüm":
                col2 = col2.split('\n')
                col2 = col2[0]
                col2 += "\n" + " ".join([i.lstrip(' ') for i in col2[1:]])
            elif col1 == "Etki":
                col2 = col2.split('\n')
                col2 = " ".join([i.lstrip(' ') for i in col2])

            row[1].text = col2
            row[1].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        try:
            self.doc.save(self.document_location)
            return "OK"
        except Exception as e:
            return e
